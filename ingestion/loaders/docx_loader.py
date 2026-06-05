"""
ingestion/loaders/docx_loader.py
"""

from __future__ import annotations

import structlog
from pathlib import Path

from domain.entities import BoundingBox, DocumentFormat, RawDocument, RawPage
from domain.exceptions import DocumentLoadError
from ingestion.loaders.base import LoaderBase

logger = structlog.get_logger(__name__)


class DocxLoader(LoaderBase):
    def load(self, path: Path, doc_id: str) -> RawDocument:
        try:
            from docx import Document
        except ImportError as exc:
            raise DocumentLoadError(doc_id, "python-docx not installed") from exc

        if not path.exists():
            raise DocumentLoadError(doc_id, f"File not found: {path}")

        try:
            doc = Document(str(path))
        except Exception as exc:
            raise DocumentLoadError(doc_id, f"Cannot open DOCX: {exc}") from exc

        # Agrupa parágrafos em "páginas" lógicas de ~50 parágrafos
        all_paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        page_size = 50
        pages: list[RawPage] = []

        for i in range(0, max(1, len(all_paragraphs)), page_size):
            chunk = all_paragraphs[i : i + page_size]
            text = "\n\n".join(chunk)
            pages.append(RawPage(
                page_number=len(pages) + 1,
                raw_text=text,
                extraction_method="python_docx",
            ))

        if not pages:
            pages.append(RawPage(page_number=1, raw_text="", extraction_method="python_docx"))

        logger.info("DOCX loaded", extra={"doc_id": doc_id, "pages": len(pages)})

        return RawDocument(
            doc_id=doc_id,
            source_path=str(path),
            detected_format=DocumentFormat.DOCX,
            detected_encoding="utf-8",
            detected_language="pt",
            is_scanned=False,
            pages=pages,
            total_pages=len(pages),
        )
